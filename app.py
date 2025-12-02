import streamlit as st
from PIL import Image
import io
import base64
import pandas as pd
from docx import Document
from docx.shared import Inches
# Importa la librería de Vertex AI
import vertexai
from vertexai.generative_models import GenerativeModel, Part, Image as VertexImage, GenerationConfig
import json
import random # Necesario para la clave aleatoria
from google.cloud import storage
import os

# --- IMPORTACIÓN CLAVE ---
# Importamos las TRES funciones que necesitamos
try:
    from graficos_plugins import crear_grafico, generar_grafico_desde_texto, build_visual_json_with_llm
    GRAFICOS_DISPONIBLES = True
except ImportError:
    st.error("Advertencia: No se encontró el archivo 'graficos_plugins.py'. La previsualización de gráficos no funcionará.")
    GRAFICOS_DISPONIBLES = False
    # Definir funciones placeholder si falla la importación
    def crear_grafico(*args, **kwargs):
        return None
    def generar_grafico_desde_texto(*args, **kwargs):
        return None, None
    def build_visual_json_with_llm(*args, **kwargs):
        return None
# --- Configuración de Google Cloud (hacer al inicio) ---
# DESCOMENTA Y RELLENA ESTAS LÍNEAS:
GCP_PROJECT = "espejazos"  # Escribe aquí el ID de tu proyecto de Google Cloud
GCP_LOCATION = "us-central1"           # O la región que estés usando (ej. us-east1)
vertexai.init(project=GCP_PROJECT, location=GCP_LOCATION)

# --- 1. FUNCIÓN DEL GENERADOR (CORREGIDA PARA TEXTO NATURAL PURO) ---
def generar_item_llm(imagen_cargada, taxonomia_dict, contexto_adicional, feedback_auditor=""):
    """
    GENERADOR: Genera el ítem, pidiendo descripciones de gráficos en LENGUAJE NATURAL PURO.
    """
    
    # --- Configuración del Modelo ---
    model = GenerativeModel("gemini-2.5-flash-lite") 
    
    # --- Procesamiento de Imagen ---
    img_pil = Image.open(imagen_cargada)
    buffered = io.BytesIO()
    img_pil.save(buffered, format="PNG")
    img_bytes = buffered.getvalue()
    vertex_img = VertexImage.from_bytes(img_bytes)

    # --- Preparación de variables del Prompt ---
    taxonomia_texto = "\n".join([f"* {k}: {v}" for k, v in taxonomia_dict.items()])
    clave_aleatoria = random.choice(['A', 'B', 'C', 'D'])

    seccion_feedback = ""
    if feedback_auditor:
        seccion_feedback = f"""
        --- RETROALIMENTACIÓN DE AUDITORÍA (Error a corregir) ---
        El intento anterior fue rechazado. DEBES corregir los siguientes errores:
        {feedback_auditor}
        --- VUELVE A GENERAR EL ÍTEM CORRIGIENDO ESTO ---
        """

    # --- 4. Diseño del Prompt (Generador) - ¡EJEMPLOS CORREGIDOS! ---
    prompt_texto = f"""
    Eres un  experto en evaluación educativa, con especialización en el diseño de ítems de lectura crítica para pruebas estandarizadas de alto impacto, como la prueba Saber 11 en Colombia. Tu misión es crear una pregunta espejo que sea un clon psicométrico de la pregunta original. Esto significa que, aunque se aplique a un texto nuevo, debe evaluar exactamente la misma habilidad, con el mismo formato y nivel de dificultad, garantizando que ambos ítems sean funcionalmente equivalentes.

    {seccion_feedback}

    **Shell Cognitivo (Pregunta Original):**
    Analiza la estructura lógica y la "Tarea Cognitiva" de la pregunta en la IMAGEN ADJUNTA.
    - Si la pregunta original usa una tabla o gráfico, tu ítem espejo también debería usar uno.
    - **¡IMPORTANTE!** Si las *opciones de respuesta* en la imagen original son gráficas o tablas, debes replicar esa estructura para las opciones del ítem espejo.
    2. Análisis de la pregunta modelo

    Identifica la habilidad cognitiva (p. ej., inferencia, comprensión literal, vocabulario).
    Observa el formato (cita breve, expresión subrayada, pregunta abierta, etc.).
    Revisa el tipo de distractores (antónimos, conceptos afines, distractores temáticos).
      
    **¡INSTRUCCIÓN CRÍTICA DE SIMILITUD!**
        1.  **NO CAMBIES LA ESTRUCTURA**: Si la pregunta usa una tabla, tu ítem espejo debe usar una tabla con la MISMA ESTRUCTURA (mismas columnas y filas).
        2.  **DEBES CAMBIAR**:
            - Los **valores numéricos** 
            - Los **nombres ficticios**
            - Los **contextos**
        3. Creación de la pregunta espejo
        Sobre el texto nuevo, elabora una pregunta que:
        Evalúe la misma competencia y evidencia, con igual nivel de dificultad.
        Repita el formato estructural.
        Garantice una respuesta correcta única y clara; los distractores deben ser plausibles, pero inequívocamente incorrectos.
        Utilice vocabulario y complejidad adecuados para estudiantes de grado 11.

    **Taxonomía Requerida (Tu Guía):**
    {taxonomia_texto}
    
    **Contexto Adicional del Usuario (Tema del ítem nuevo):**
    {contexto_adicional}

    --- ANÁLISIS COGNITIVO OBLIGATORIO (Tu paso 1) ---
    Basado en la taxonomía (Evidencia, Afirmación, Competencia), define la Tarea Cognitiva exacta que el ítem espejo debe evaluar.
    
    --- CONSTRUCCIÓN DEL ÍTEM (Tu paso 2) ---
    Basado en tu análisis, construye el ítem.
    - ENUNCIADO: Debe ser claro y **NO** usar jerarquías ("más", "mejor", "principalmente").
    - CLAVE: La respuesta correcta DEBE ser la opción **{clave_aleatoria}**.
    - DISTRACTORES: Plausibles, basados en errores comunes de la Tarea Cognitiva. Deben tener la redacción "El estudiante podria escoger la opción XX porque... Sin embargo esto es incorrecto porque...
    - DIFERENCIAS CON EL ITEM INICIAL: *CRITICO* NO se puede usar ninguno de los valores numéricos del ítem inicial. Deben ser totalmente diferentes.
    
    
    --- INSTRUCCIONES DE SALIDA PARA GRÁFICO (ENUNCIADO Y OPCIONES) ---
    ¡INSTRUCCIÓN CRÍTICA! Para los gráficos, NO debes generar el JSON.
    En su lugar, proporciona una descripción detallada en LENGUAJE NATURAL de lo que el gráfico debe mostrar.
    
    Si el elemento (enunciado u opción) NO necesita un gráfico, usa "NO" y "N/A".
    Si SÍ necesita un gráfico, usa "SÍ" y escribe la descripción.
    
    Ejemplo de descripción: "Una tabla de 3 columnas y 2 filas. Las columnas son 'País', 'Capital', 'Población'. La primera fila es 'Colombia', 'Bogotá', '8M'. La segunda es 'Argentina', 'Buenos Aires', '3M'."
    Otro ejemplo: "Un gráfico de barras verticales simple con 3 barras. El eje X tiene las etiquetas 'A', 'B', 'C'. El eje Y (valores) tiene '10', '20', '15'."

    --- FORMATO DE SALIDA OBLIGATORIO (JSON VÁLIDO) ---
    Responde ÚNICAMENTE con el objeto JSON. No incluyas ```json.
    {{
      "pregunta_espejo": "Texto completo del enunciado/stem...",
      "clave": "{clave_aleatoria}",
      "descripcion_imagen_original": "Descripción de la imagen que el usuario subió...",
      "justificacion_clave": "Razón por la que la clave es correcta...",
      
      "grafico_necesario_enunciado": "SÍ",
      "descripcion_texto_grafico_enunciado": "Una tabla simple. La primera fila es el encabezado con 'País' y 'Capital'. La segunda fila tiene 'Colombia' y 'Bogotá'.",
      
      "opciones": {{
        "A": {{
          "texto": "Ver gráfico A",
          "grafico_necesario": "SÍ",
          "descripcion_texto_grafico": "Un gráfico de barras verticales simple. El eje X tiene dos categorías: 'X' y 'Y'. Los valores del eje Y son 5 para 'X' y 10 para 'Y'."
        }},
        "B": {{
          "texto": "Texto de la Opción B (sin gráfico)",
          "grafico_necesario": "NO",
          "descripcion_texto_grafico": "N/A"
        }},
        "C": {{
          "texto": "Texto de la Opción C",
          "grafico_necesario": "NO",
          "descripcion_texto_grafico": "N/A"
        }},
        "D": {{
          "texto": "Texto de la Opción D",
          "grafico_necesario": "NO",
          "descripcion_texto_grafico": "N/A"
        }}
      }},
      
      "justificaciones_distractores": [
        {{ "opcion": "A", "justificacion": "Justificación para A..." }},
        {{ "opcion": "B", "justificacion": "Justificación para B..." }},
        {{ "opcion": "C", "justificacion": "Justificación para C..." }},
        {{ "opcion": "D", "justificacion": "Justificación para D..." }}
      ]
    }}
    """

    config_generacion = GenerationConfig(
        response_mime_type="application/json"
    )

    try:
        # --- 1. LLAMADA A LA API ---
        response = model.generate_content(
            [vertex_img, prompt_texto], 
            generation_config=config_generacion
        )
        
        raw_text = response.text
        
        # --- 2. MEJORA: LIMPIEZA DE JSON ---
        try:
            start_index = raw_text.find('{')
            end_index = raw_text.rfind('}') + 1
            if start_index == -1 or end_index == 0:
                raise ValueError("No se encontraron los delimitadores JSON '{' o '}'.")
            json_str = raw_text[start_index:end_index]
            json.loads(json_str) 
            return json_str
        
        except (ValueError, json.JSONDecodeError) as json_e:
            st.error(f"Error al limpiar/parsear la respuesta del Generador: {json_e}")
            st.error(f"Respuesta cruda recibida (esto puede ayudar a depurar): {raw_text}")
            return None
        # --- FIN DE LA MEJORA DE LIMPIEZA ---

    except Exception as e:
        st.error(f"Error al contactar Vertex AI (Generador): {e}")
        return None

# --- 2. FUNCIÓN DEL AUDITOR (ACTUALIZADA CON LIMPIEZA DE JSON) ---
def auditar_item_llm(item_json_texto, taxonomia_dict):
    """
    AUDITOR: Audita el ítem Y la coherencia de los gráficos (enunciado Y opciones).
    """
    
    # Modelo de Gemini (corregido al que usas)
    model = GenerativeModel("gemini-2.5-flash-lite")
    taxonomia_texto = "\n".join([f"* {k}: {v}" for k, v in taxonomia_dict.items()])

    prompt_auditor = f"""
    Eres un auditor psicométrico experto y riguroso. Tu tarea es auditar el siguiente ítem (en JSON)
    contra la taxonomía y las reglas de estilo.
    
    **Taxonomía de Referencia (ObligatorIA):**
    {taxonomia_texto}

    **Ítem Generado (JSON a Auditar):**
    {item_json_texto}

    --- CRITERIOS DE AUDITORÍA (Evalúa uno por uno) ---
    1.  **Alineación con Taxonomía:** ¿El ítem evalúa CLARAMENTE la Evidencia, Afirmación y Competencia?
    2.  **Estilo del Enunciado (No Jerarquización):** ¿El enunciado usa palabras prohibidas como "más", "mejor", "principalmente"? (RECHAZO automático).
    3.  **Calidad de Distractores:** ¿Las justificaciones de los distractores explican el *error* (ej. "El estudiante podría...")?
    4.  **Clave y Opciones:** ¿Hay 4 opciones? ¿La clave coincide con una opción?
    5.  **Coherencia de Gráficos (¡ACTUALIZADO!):** - ¿Es coherente el "grafico_necesario_enunciado" con la pregunta?
        - ¿Son coherentes los "grafico_necesario" DENTRO de cada opción?
        - Si un gráfico existe, ¿es un JSON válido?

    --- FORMATO DE SALIDA OBLIGATORIO (JSON VÁLIDO) ---
    Devuelve tu auditoría como un único objeto JSON. No uses ```json.
    {{
      "criterios": [
        {{ "criterio": "1. Alineación con Taxonomía", "estado": "✅ CUMPLE" o "❌ NO CUMPLE", "comentario": "Justificación breve." }},
        {{ "criterio": "2. Estilo (No Jerarquización)", "estado": "✅ CUMPLE" o "❌ NO CUMPLE", "comentario": "Justificación breve." }},
        {{ "criterio": "3. Calidad de Distractores", "estado": "✅ CUMPLE" o "❌ NO CUMPLE", "comentario": "Justificación breve." }},
        {{ "criterio": "4. Clave y Opciones", "estado": "✅ CUMPLE" o "❌ NO CUMPLE", "comentario": "Justificación breve." }},
        {{ "criterio": "5. Coherencia de Gráficos", "estado": "✅ CUMPLE" o "❌ NO CUMPLE", "comentario": "Justificación breve." }}
      ],
      "dictamen_final": "✅ CUMPLE" o "❌ RECHAZADO",
      "observaciones_finales": "Si es RECHAZADO, explica aquí CLARAMENTE qué debe corregir el generador. (Ej: 'El enunciado usa la palabra 'principalmente'. O 'El gráfico de la opción C es SÍ pero no se proporcionó JSON.')"
    }}
    """
    
    config_generacion = GenerationConfig(
        response_mime_type="application/json"
    )

    try:
        response = model.generate_content(
            prompt_auditor, 
            generation_config=config_generacion
        )
        
        raw_text = response.text
        
        # --- INICIO DE LA MEJORA: LIMPIEZA DE JSON (AUDITOR) ---
        try:
            # Encuentra el primer { y el último } para eliminar texto extra
            start_index = raw_text.find('{')
            end_index = raw_text.rfind('}') + 1
            
            if start_index == -1 or end_index == 0:
                raise ValueError("No se encontraron los delimitadores JSON '{' o '}'.")

            # Extrae solo el JSON
            json_str = raw_text[start_index:end_index]
            
            # Valida que es un JSON antes de devolver
            json.loads(json_str) 
            return json_str
        
        except (ValueError, json.JSONDecodeError) as json_e:
            st.error(f"Error al limpiar/parsear la respuesta del Auditor: {json_e}")
            st.error(f"Respuesta cruda recibida (esto puede ayudar a depurar): {raw_text}")
            return None
        # --- FIN DE LA MEJORA DE LIMPIEZA ---

    except Exception as e:
        st.error(f"Error al contactar Vertex AI (Auditor): {e}")
        return None

# --- 3. FUNCIONES DE EXPORTACIÓN (ACTUALIZADAS) ---

# --- 3. FUNCIONES DE EXPORTACIÓN (ACTUALIZADAS) ---

def reemplazar_texto_en_doc(doc, reemplazos):
    """
    Recorre todos los párrafos y tablas en un documento y reemplaza los placeholders.
    """
    for p in doc.paragraphs:
        for clave, valor in reemplazos.items():
            if clave in p.text:
                p.text = p.text.replace(clave, valor)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for clave, valor in reemplazos.items():
                        if clave in p.text:
                            p.text = p.text.replace(clave, valor)
    return doc

# --- 3. FUNCIONES DE EXPORTACIÓN (EXCEL REESCRITO CON AFIRMACIÓN) ---

def crear_excel(datos_generados, taxonomia_seleccionada, oportunidad_mejora):
    """
    Crea un Excel en formato HORIZONTAL (una fila por ítem) con las columnas
    específicas solicitadas.
    """
    
    # 1. Mapear las justificaciones a un diccionario para fácil acceso
    justificaciones = datos_generados.get("justificaciones_distractores", [])
    justifs_map = {j.get('opcion'): j.get('justificacion') for j in justificaciones}
    
    # 2. Crear el diccionario de datos para la única fila
    data_dict = {
        # --- Columnas de Taxonomía ---
        "Área": taxonomia_seleccionada.get("Área", "N/A"),
        "RESPONSABLE": "IA ESPEJAZOS",
        "COMPONENTE": taxonomia_seleccionada.get("Componente_Estructura", "N/A"),
        "Competencia": taxonomia_seleccionada.get("Competencia", "N/A"),
        "Afirmación": taxonomia_seleccionada.get("Afirmación", "N/A"), # <-- AÑADIDA DE VUELTA
        "Evidencia": taxonomia_seleccionada.get("Evidencia", "N/A"),
        "Temática": taxonomia_seleccionada.get("Ref. Temática", "N/A"),
        "Nivel (curso)": taxonomia_seleccionada.get("Grado", "N/A"),
        
        # --- Columnas de Metadatos (Fijas) ---
        "PASTILLA": "NA",
        "Dificultad estimada": "NA",
        "Estándar": "NA",
        "Estado": "Espejo",
        "Número en el PDF": "NA",
        "ID ÍTEM": "NA",
        "ID CONTEXTO": "NA",

        # --- Columnas de Contenido del Ítem ---
        "Guía (Primeras palabras del ítem)": datos_generados.get("pregunta_espejo", "N/A"),
        "Oportunidad de mejora": oportunidad_mejora,
        "Justificación de la respuesta A": justifs_map.get("A", "N/A"),
        "Justificación de la respuesta B": justifs_map.get("B", "N/A"),
        "Justificación de la respuesta C": justifs_map.get("C", "N/A"),
        "Justificación de la respuesta D": justifs_map.get("D", "N/A"),
        "Clave": datos_generados.get("clave", "N/A")
    }

    # 3. Crear el DataFrame
    df = pd.DataFrame([data_dict])
    
    # 4. Forzar el orden de columnas exacto que pediste
    columnas_ordenadas = [
        "Área",
        "RESPONSABLE",
        "COMPONENTE",
        "Competencia",
        "Afirmación", # <-- AÑADIDA DE VUELTA
        "Evidencia",
        "PASTILLA",
        "Temática",
        "Dificultad estimada",
        "Estándar",
        "Estado",
        "Nivel (curso)",
        "Número en el PDF",
        "ID ÍTEM",
        "ID CONTEXTO",
        "Guía (Primeras palabras del ítem)",
        "Oportunidad de mejora",
        "Justificación de la respuesta A",
        "Justificación de la respuesta B",
        "Justificación de la respuesta C",
        "Justificación de la respuesta D",
        "Clave"
    ]
    
    # Filtra solo las columnas que existen en el df para evitar errores
    columnas_finales = [col for col in columnas_ordenadas if col in df.columns]
    df = df[columnas_finales]

    # 5. Guardar en el buffer de Excel
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Item Generado')
    return output.getvalue()

def crear_word(datos_editados, taxonomia_seleccionada, oportunidad_mejora):
    """
    Genera un documento Word rellenando una plantilla desde GCS.
    """
    try:
        # 1. Descargar la plantilla desde GCS
        bucket_name = "bucket-espejos"
        template_name = "formato_limpio.docx"
        storage_client = storage.Client(project=GCP_PROJECT)
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(template_name)
        
        if not blob.exists():
            st.error(f"Error: La plantilla '{template_name}' no se encontró en el bucket '{bucket_name}'.")
            return None
            
        doc_buffer = io.BytesIO(blob.download_as_bytes())
        doc = Document(doc_buffer)
        
        # 2. Oportunidad de mejora (Ahora se recibe como argumento)
        # (La llamada a la IA se eliminó de aquí)
        
        # 3. Preparar los distractores
        clave = datos_editados.get("clave", "")
        distractores_texto = []
        for just in datos_editados.get("justificaciones_distractores", []):
            opcion = just.get("opcion")
            if opcion and opcion != clave:
                distractores_texto.append(f"Opción {opcion}: {just.get('justificacion', 'N/A')}")
        analisis_distractores = "\n".join(distractores_texto)

        # 4. Preparar las instrucciones de gráficos (convierte JSON a string)
        def get_grafico_json(data):
            if not data or data == "[]" or data == []:
                return "N/A"
            # Usamos ensure_ascii=False para que no escape tildes (ej. \u00f3)
            return json.dumps(data, ensure_ascii=False, indent=2)

        inst_enunciado = get_grafico_json(datos_editados.get("descripcion_grafico_enunciado", []))
        inst_a = get_grafico_json(datos_editados.get("opciones", {}).get("A", {}).get("descripcion_grafico", []))
        inst_b = get_grafico_json(datos_editados.get("opciones", {}).get("B", {}).get("descripcion_grafico", []))
        inst_c = get_grafico_json(datos_editados.get("opciones", {}).get("C", {}).get("descripcion_grafico", []))
        inst_d = get_grafico_json(datos_editados.get("opciones", {}).get("D", {}).get("descripcion_grafico", []))

        # 5. Definir todos los reemplazos (¡TODOS CON str()!)
        reemplazos = {
            "{{ItemPruebaId}}": str(taxonomia_seleccionada.get("Área", "N/A")),
            "{{ItemGradoId}}": str(taxonomia_seleccionada.get("Grado", "N/A")), 
            "{{CompetenciaNombre}}": str(taxonomia_seleccionada.get("Competencia", "N/A")),
            "{{ComponenteNombre}}": str(taxonomia_seleccionada.get("Componente_Estructura", "N/A")),
            "{{AfirmacionNombre}}": str(taxonomia_seleccionada.get("Afirmación", "N/A")),
            "{{EvidenciaNombre}}": str(taxonomia_seleccionada.get("Evidencia", "N/A")),
            "{{ItemContexto}}": "", 
            "{{ItemEnunciado}}": str(datos_editados.get("pregunta_espejo", "N/A")),
            "{{Opción A}}": str(datos_editados.get("opciones", {}).get("A", {}).get("texto", "N/A")),
            "{{Opción B}}": str(datos_editados.get("opciones", {}).get("B", {}).get("texto", "N/A")),
            "{{Opción C}}": str(datos_editados.get("opciones", {}).get("C", {}).get("texto", "N/A")),
            "{{Opción D}}": str(datos_editados.get("opciones", {}).get("D", {}).get("texto", "N/A")),
            "{{  Clave}}": str(clave),
            "{{Justificacion_Correcta}}": str(datos_editados.get("justificacion_clave", "N/A")),
            "{{Analisis_Distractores}}": str(analisis_distractores),
            "{{Instrucciones_enuncuado}}": str(inst_enunciado),
            "{{Instrucciones_A}}": str(inst_a),
            "{{Instrucciones_B}}": str(inst_b),
            "{{Instrucciones_C}}": str(inst_c),
            "{{Instrucciones_D}}": str(inst_d),
            "Enunciado: {{Instrucciones_D}}": str(inst_d), # Por si acaso
            "{{Oportunidad_mejora}}": str(oportunidad_mejora)
        }

        # 6. Ejecutar los reemplazos
        doc = reemplazar_texto_en_doc(doc, reemplazos)

        # 7. Guardar el documento final en un nuevo buffer
        final_buffer = io.BytesIO()
        doc.save(final_buffer)
        final_buffer.seek(0)
        return final_buffer

    except Exception as e:
        st.error(f"Error al crear el documento Word: {e}")
        return None


# --- NUEVA FUNCIÓN PARA LEER EXCEL DESDE GCS ---
@st.cache_data
def leer_excel_desde_gcs(bucket_name, file_path):
    """
    Lee un archivo Excel (con todas sus hojas) directamente desde GCS.
    """
    try:
        storage_client = storage.Client(project=GCP_PROJECT) # Usa el proyecto ya inicializado
        bucket = storage_client.bucket(bucket_name)
        blob = bucket.blob(file_path)
        
        if not blob.exists():
            st.error(f"Error: El archivo '{file_path}' no se encontró en el bucket '{bucket_name}'.")
            return None
            
        file_bytes = blob.download_as_bytes()
        # Cargar todas las hojas del Excel
        data = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
        st.success(f"Éxito: Cargado '{file_path}' desde GCS.")
        return data
    except Exception as e:
        st.error(f"Error al leer Excel desde GCS: {e}")
        st.info("Asegúrate de que la cuenta de servicio de Streamlit tenga permisos de 'Storage Object Viewer' en el bucket 'bucket-espejos'.")
        return None
# --- FIN DE LA NUEVA FUNCIÓN ---


# --- 2b. NUEVA FUNCIÓN: GENERADOR DE OPORTUNIDAD DE MEJORA ---
def generar_oportunidad_mejora_llm(taxonomia_data, justificacion_clave):
    """
    Genera una breve recomendación académica basada en la habilidad evaluada.
    """
    try:
        model = GenerativeModel("gemini-2.5-flash-lite")
        
        # Extraemos los datos clave para el prompt
        evidencia = taxonomia_data.get("Evidencia", "la habilidad evaluada")
        competencia = taxonomia_data.get("Competencia", "la competencia general")
        
        prompt = f"""
        Eres un tutor académico experto. Tu tarea es escribir una breve recomendación (1-2 frases)
        para un estudiante o profesor.
        
        HABILIDAD EVALUADA (Evidencia): {evidencia}
        COMPETENCIA: {competencia}
        JUSTIFICACIÓN DE LA RESPUESTA CORRECTA: {justificacion_clave}

        Basado en esta información, escribe una recomendación fácil de aplicar durante clase.
        Habla en tercera persona y manten un tono formal pero sencillo
        NO uses más de 50 palabras.
        """
        
        response = model.generate_content(prompt)
        return response.text.strip()
    
    except Exception as e:
        print(f"Error al generar oportunidad de mejora: {e}")
        return "Para mejorar en esta habilidad, repasa los conceptos clave de la competencia y practica con ejercicios similares."



# --- 4. INTERFAZ DE STREAMLIT (UI) ---

st.set_page_config(layout="wide")
st.title("🤖 Generador de Ítems espejo (con Auditoría de IA)")

# --- Columnas para la entrada ---
col1, col2 = st.columns(2)

with col1:
    st.header("1. Cargar Ítem Original")
    imagen_subida = st.file_uploader(
        "Sube el pantallazo de la pregunta", 
        type=["png", "jpg", "jpeg"]
    )
    
    if imagen_subida:
        st.image(imagen_subida, caption="Ítem cargado", use_container_width=True)

# --- COLUMNA 2 (Lógica de Filtros Bifurcada y CORREGIDA) ---
with col2:
    st.header("2. Configurar Generación")
    
    # --- MODIFICACIÓN: Cargar Excel desde GCS ---
    bucket_name = "bucket-espejos"
    excel_file_path = "Estructura privados1.xlsx"
    
    # Intentamos cargar los datos desde GCS
    data = leer_excel_desde_gcs(bucket_name, excel_file_path)
    # --- FIN DE MODIFICACIÓN ---
    
    grado_sel, area_sel, comp1_sel, comp2_sel, ref_sel, competen_sel, afirm_sel, evid_sel = (None,) * 8
    
    if data is not None: # Cambiamos 'excel_file is not None' por 'data is not None'
        try:
            if 'df1' not in st.session_state or 'df2' not in st.session_state:
                # data ya está cargado por la función de GCS
                sheet_names = list(data.keys())
                if len(sheet_names) < 2:
                    st.error("Error: El archivo Excel debe tener al menos dos hojas.")
                    data = None # Anulamos data para detener la ejecución
                else:
                    st.session_state.df1 = data[sheet_names[0]]
                    st.session_state.df2 = data[sheet_names[1]]
                    # La función leer_excel_desde_gcs ya muestra el st.success
            
            if 'df1' in st.session_state:
                df1 = st.session_state.df1
                df2 = st.session_state.df2

                # --- Filtros Comunes ---
                grados = df1['Grado'].unique()
                grado_sel = st.selectbox("Grado", options=grados)
                
                df_grado_h1 = df1[df1['Grado'] == grado_sel]
                areas = df_grado_h1['Área'].unique() # Con tilde
                area_sel = st.selectbox("Área", options=areas) # Con tilde

                # --- Cascada 1: (Hoja 1 - Estructura) ---
                st.subheader("Taxonomía (Hoja 1 - Estructura)")
                df_area_h1 = df_grado_h1[df_grado_h1['Área'] == area_sel]
                componentes1 = df_area_h1['Componente1'].unique() 
                comp1_sel = st.selectbox("Componente (Estructura)", options=componentes1) 

                df_comp1 = df_area_h1[df_area_h1['Componente1'] == comp1_sel]
                competencias = df_comp1['Competencia'].unique()
                competen_sel = st.selectbox("Competencia", options=competencias)

                df_competencia = df_comp1[df_comp1['Competencia'] == competen_sel]
                
                if area_sel == 'Ciencias Naturales': 
                    df_afirmacion_base = df_competencia[df_competencia['Componente1'] == comp1_sel]
                else:
                    df_afirmacion_base = df_competencia
                    
                afirmaciones = df_afirmacion_base['Afirmación'].unique()
                afirm_sel = st.selectbox("Afirmación", options=afirmaciones)

                df_afirmacion = df_afirmacion_base[df_afirmacion_base['Afirmación'] == afirm_sel]
                evidencias = df_afirmacion['Evidencia'].unique()
                evid_sel = st.selectbox("Evidencia", options=evidencias)

                # --- Cascada 2: (Hoja 2 - Temática) ---
                st.subheader("Taxonomía (Hoja 2 - Temática)")
                df_area_h2 = df2[
                    (df2['Grado'] == grado_sel) & 
                    (df2['Área'] == area_sel) # Con tilde
                ]
                componentes2 = df_area_h2['Componente2'].unique()
                comp2_sel = st.selectbox("Componente (Temática)", options=componentes2)

                df_comp2 = df_area_h2[df_area_h2['Componente2'] == comp2_sel]
                
                refs = df_comp2['Ref. Temática'].unique() if not df_comp2.empty else ["N/A"] # Con tilde y espacio
                ref_sel = st.selectbox("Ref. Temática", options=refs) # Con tilde y espacio

        except KeyError as e:
            st.error(f"Error de Columna: No se encontró la columna {e}. Revisa las tildes/mayúsculas.")
            if 'df1' in st.session_state: st.error(f"Columnas H1: {list(st.session_state.df1.columns)}")
            if 'df2' in st.session_state: st.error(f"Columnas H2: {list(st.session_state.df2.columns)}")
            data = None # Anulamos data para detener la ejecución
        except Exception as e:
            st.error(f"Error inesperado al procesar el Excel: {e}")
            data = None # Anulamos data para detener la ejecución
    
    info_adicional = st.text_area(
        "Contexto Adicional (Tema para el ítem)",
        height=150,
        placeholder="Ej: 'Usar el tema de la fotosíntesis', 'Basarse en la Revolución Francesa'"
    )

# --- 5. LÓGICA DEL BOTÓN (Bucle Generador-Auditor) ---
st.divider()
if st.button("🚀 Generar Ítem Espejo (con Auditoría)", use_container_width=True, type="primary"):
    
    if imagen_subida is None:
        st.warning("Por favor, sube una imagen primero.")
    elif data is None: # <-- ¡CORRECCIÓN CLAVE!
        st.warning("El archivo Excel de taxonomía no se pudo cargar desde GCS. Revisa los errores en la Columna 2.")
    elif evid_sel is None or ref_sel is None:
        st.warning("Completa toda la selección de taxonomía.")
    else:
        taxonomia_seleccionada = {
            "Grado": grado_sel,
            "Área": area_sel,
            "Componente_Estructura": comp1_sel, # Nombre corregido
            "Componente_Tematica": comp2_sel,  # Nombre corregido
            "Ref. Temática": ref_sel,
            "Competencia": competen_sel,
            "Afirmación": afirm_sel,
            "Evidencia": evid_sel
        }
        st.session_state['taxonomia_actual'] = taxonomia_seleccionada
        
        max_intentos = 3
        intento_actual = 0
        feedback_auditor = ""
        item_final_json = None

        with st.status("Iniciando proceso...", expanded=True) as status:
            while intento_actual < max_intentos:
                intento_actual += 1
                
                status.update(label=f"Intento {intento_actual}/{max_intentos}: Generando ítem...")
                item_json_str = generar_item_llm(
                    imagen_subida, 
                    taxonomia_seleccionada,
                    info_adicional,
                    feedback_auditor 
                )
                
                if item_json_str is None:
                    status.update(label=f"Error en la generación (Intento {intento_actual}).", state="error")
                    continue 

                status.update(label=f"Intento {intento_actual}/{max_intentos}: Auditando ítem...")
                audit_json_str = auditar_item_llm(item_json_str, taxonomia_seleccionada)

                if audit_json_str is None:
                    status.update(label=f"Error en la auditoría (Intento {intento_actual}).", state="error")
                    continue 

                try:
                    # --- FIX: Asegurarse de parsear la respuesta del auditor ---
                    audit_data = json.loads(audit_json_str)
                    
                    if audit_data.get("dictamen_final") == "✅ CUMPLE":
                        status.update(label="¡Auditoría Aprobada!", state="complete")
                        item_final_json = item_json_str
                        break 
                    else:
                        feedback_auditor = audit_data.get("observaciones_finales", "Rechazado sin observaciones.")
                        status.update(label=f"Intento {intento_actual} Rechazado. Preparando re-intento...")
                        st.expander(f"Detalles del Rechazo (Intento {intento_actual})").json(audit_data)
                
                except json.JSONDecodeError:
                    st.error(f"Error al leer respuesta JSON del auditor: {audit_json_str}")
                    feedback_auditor = "La respuesta del auditor no fue un JSON válido."

            if item_final_json is None:
                status.update(label=f"No se pudo generar un ítem de alta calidad después de {max_intentos} intentos.", state="error")
                st.error(f"Último feedback del auditor: {feedback_auditor}")
            
        if item_final_json:
            st.success("¡Ítem generado y auditado con éxito! Puedes editarlo abajo.")
            try:
                # --- FIX: Asegurarse de parsear la respuesta del generador ---
                datos_obj = json.loads(item_final_json)
                st.session_state['resultado_json_obj'] = datos_obj
                
                # --- LÓGICA DE INICIALIZACIÓN (ACTUALIZADA para nuevo JSON con TEXTO) ---
                st.session_state.editable_pregunta = datos_obj.get("pregunta_espejo", "")
                st.session_state.editable_clave = datos_obj.get("clave", "")
                st.session_state.editable_just_clave = datos_obj.get("justificacion_clave", "")
    
                # Gráfico del Enunciado
                st.session_state.editable_grafico_nec_enunciado = datos_obj.get("grafico_necesario_enunciado", "NO")
                # NUEVO: Guardamos la descripción de TEXTO
                st.session_state.editable_grafico_texto_enunciado = datos_obj.get("descripcion_texto_grafico_enunciado", "N/A")
                # INICIALIZIAMOS EL JSON COMO VACÍO
                st.session_state.editable_grafico_json_enunciado = "[]"
    
                # Opciones (A, B, C, D)
                opciones = datos_obj.get("opciones", {})
                for letra in ["A", "B", "C", "D"]:
                    opcion_obj = opciones.get(letra, {}) 
                    
                    st.session_state[f"editable_opcion_{letra.lower()}_texto"] = opcion_obj.get("texto", "")
                    st.session_state[f"editable_opcion_{letra.lower()}_grafico_nec"] = opcion_obj.get("grafico_necesario", "NO")
                    # NUEVO: Guardamos la descripción de TEXTO
                    st.session_state[f"editable_opcion_{letra.lower()}_grafico_texto"] = opcion_obj.get("descripcion_texto_grafico", "N/A")
                    # INICIALIZIAMOS EL JSON COMO VACÍO
                    st.session_state[f"editable_opcion_{letra.lower()}_grafico_json"] = "[]"
    
                # Justificaciones
                justifs_list = datos_obj.get("justificaciones_distractores", [])
                justifs_map = {j.get('opcion'): j.get('justificacion') for j in justifs_list}
                st.session_state.editable_just_a = justifs_map.get("A", "N/A")
                st.session_state.editable_just_b = justifs_map.get("B", "N/A")
                st.session_state.editable_just_c = justifs_map.get("C", "N/A")
                st.session_state.editable_just_d = justifs_map.get("D", "N/A")
                
                st.session_state.show_editor = True
                
            except json.JSONDecodeError:
                st.error(f"Error al parsear el JSON final: {item_final_json}")
                st.session_state.show_editor = False

# --- 6. EDITOR DE ÍTEMS Y DESCARGA (LÓGICA DE BOTONES SEPARADA) ---
if 'show_editor' in st.session_state and st.session_state.show_editor:
    st.divider()
    st.header("3. Edita el Ítem Generado")
    
    # --- ENUNCIADO Y GRÁFICO DEL ENUNCIADO ---
    st.subheader("Enunciado")
    st.text_area("Texto del Enunciado", key="editable_pregunta", height=150)
    st.selectbox(
        "¿Enunciado necesita un gráfico/tabla?", 
        options=["NO", "SÍ"], 
        key="editable_grafico_nec_enunciado"
    )
    
    if st.session_state.editable_grafico_nec_enunciado == "SÍ":
        st.text_area(
            "Descripción de Texto (Generada por IA)", 
            key="editable_grafico_texto_enunciado", 
            height=100
        )
        
        # --- Botón 1: Generar JSON ---
        if st.button("Generar JSON desde Texto (Enunciado) 🤖", key="btn_gen_json_enunciado"):
            if GRAFICOS_DISPONIBLES:
                with st.spinner("Llamando a IA de plugins para generar JSON..."):
                    texto_desc = st.session_state.editable_grafico_texto_enunciado
                    # LLAMADA A LA IA QUE SOLO HACE JSON
                    spec = build_visual_json_with_llm(texto_desc)
                    
                    if spec:
                        # GUARDAMOS EL JSON CORRECTO EN EL EDITOR
                        st.session_state.editable_grafico_json_enunciado = json.dumps([spec], indent=2)
                        st.session_state['img_buffer_enunciado'] = None # Limpiamos la imagen anterior
                        st.success("¡JSON generado! Ahora puedes editarlo o generar el gráfico.")
                    else:
                        st.error("La IA de plugins no pudo generar un JSON con esa descripción.")
            else:
                st.warning("El módulo 'graficos_plugins.py' no está disponible.")

        st.text_area(
            "Datos del Gráfico (JSON) - (Editable)", 
            key="editable_grafico_json_enunciado", 
            height=150
        )
        
        # --- Botón 2: Renderizar Gráfico desde JSON ---
        if st.button("Generar/Actualizar Gráfico desde JSON (Enunciado) 🖼️", key="btn_render_enunciado"):
            if GRAFICOS_DISPONIBLES:
                try:
                    json_data = json.loads(st.session_state.editable_grafico_json_enunciado)
                    if json_data and isinstance(json_data, list):
                        spec = json_data[0]
                        # LLAMADA AL RENDERIZADOR SIMPLE
                        buffer_imagen = crear_grafico(
                            tipo_grafico=spec.get("tipo_elemento"),
                            datos=spec.get("datos", {}),
                            configuracion=spec.get("configuracion", {})
                        )
                        if buffer_imagen:
                            st.session_state['img_buffer_enunciado'] = buffer_imagen
                            st.success("Previsualización actualizada desde JSON.")
                        else:
                            st.session_state['img_buffer_enunciado'] = None
                            st.error("No se pudo renderizar el gráfico desde el JSON. Revisa el formato.")
                except Exception as e:
                    st.session_state['img_buffer_enunciado'] = None
                    st.error(f"Error al renderizar JSON: {e}")
        
        # Mostramos la imagen si existe en el estado
        if 'img_buffer_enunciado' in st.session_state and st.session_state.img_buffer_enunciado:
            st.image(st.session_state.img_buffer_enunciado, caption="Previsualización generada")


    # --- OPCIONES Y SUS GRÁFICOS ---
    st.subheader("Opciones")
    
    for letra in ["A", "B", "C", "D"]:
        st.markdown(f"--- \n**Opción {letra}**")
        st.text_input(f"Texto Opción {letra}", key=f"editable_opcion_{letra.lower()}_texto")
        st.selectbox(
            f"¿Gráfico en Opción {letra}?", 
            options=["NO", "SÍ"], 
            key=f"editable_opcion_{letra.lower()}_grafico_nec"
        )
        
        if st.session_state[f"editable_opcion_{letra.lower()}_grafico_nec"] == "SÍ":
            st.text_area(
                f"Descripción de Texto (Opción {letra})", 
                key=f"editable_opcion_{letra.lower()}_grafico_texto", 
                height=100
            )
            
            # --- Botón 1: Generar JSON ---
            if st.button(f"Generar JSON desde Texto (Opción {letra}) 🤖", key=f"btn_gen_json_op_{letra}"):
                if GRAFICOS_DISPONIBLES:
                    with st.spinner(f"Llamando a IA de plugins para generar JSON (Opción {letra})..."):
                        texto_desc = st.session_state[f"editable_opcion_{letra.lower()}_grafico_texto"]
                        # LLAMADA A LA IA QUE SOLO HACE JSON
                        spec = build_visual_json_with_llm(texto_desc)
                        
                        if spec:
                            # GUARDAMOS EL JSON CORRECTO EN EL EDITOR
                            st.session_state[f"editable_opcion_{letra.lower()}_grafico_json"] = json.dumps([spec], indent=2)
                            st.session_state[f'img_buffer_op_{letra}'] = None # Limpiamos la imagen anterior
                            st.success(f"¡JSON generado para Opción {letra}!")
                        else:
                            st.session_state[f'img_buffer_op_{letra}'] = None
                            st.error(f"La IA de plugins no pudo generar un JSON (Opción {letra}).")
                else:
                    st.warning("El módulo 'graficos_plugins.py' no está disponible.")

            st.text_area(
                f"Datos Gráfico Opción {letra} (JSON) - (Editable)", 
                key=f"editable_opcion_{letra.lower()}_grafico_json", 
                height=150
            )
            
            # --- Botón 2: Renderizar Gráfico desde JSON ---
            if st.button(f"Generar/Actualizar Gráfico desde JSON (Opción {letra}) 🖼️", key=f"btn_render_op_{letra}"):
                if GRAFICOS_DISPONIBLES:
                    try:
                        json_data = json.loads(st.session_state[f"editable_opcion_{letra.lower()}_grafico_json"])
                        if json_data and isinstance(json_data, list):
                            spec = json_data[0]
                            # LLAMADA AL RENDERIZADOR SIMPLE
                            buffer_imagen = crear_grafico(
                                tipo_grafico=spec.get("tipo_elemento"),
                                datos=spec.get("datos", {}),
                                configuracion=spec.get("configuracion", {})
                            )
                            if buffer_imagen:
                                st.session_state[f'img_buffer_op_{letra}'] = buffer_imagen
                                st.success(f"Previsualización actualizada desde JSON (Opción {letra}).")
                            else:
                                st.session_state[f'img_buffer_op_{letra}'] = None
                                st.error(f"No se pudo renderizar el gráfico desde el JSON (Opción {letra}).")
                    except Exception as e:
                        st.session_state[f'img_buffer_op_{letra}'] = None
                        st.error(f"Error al renderizar JSON: {e}")

            # Mostramos la imagen si existe en el estado
            if f'img_buffer_op_{letra}' in st.session_state and st.session_state[f'img_buffer_op_{letra}']:
                st.image(st.session_state[f'img_buffer_op_{letra}'], caption=f"Previsualización Opción {letra}")

    st.subheader("Clave")
    st.text_input("Clave (Respuesta Correcta)", key="editable_clave")

    st.subheader("Justificaciones")
    st.text_area("Justificación Clave", key="editable_just_clave", height=100)
    st.text_area("Justificación A", key="editable_just_a", height=100)
    st.text_area("Justificación B", key="editable_just_b", height=100)
    st.text_area("Justificación C", key="editable_just_c", height=100)
    st.text_area("Justificación D", key="editable_just_d", height=100)

    # --- SECCIÓN DE DESCARGA (SIN CAMBIOS EN LA LÓGICA, USA EL JSON) ---
    # --- SECCIÓN DE DESCARGA (ACTUALIZADA PARA NUEVO EXCEL) ---
    st.divider()
    st.header("4. Descargar Resultados")
    
    # --- LÓGICA DE RE-ENSAMBLE (ACTUALIZADA) ---
    datos_editados = {
        "pregunta_espejo": st.session_state.editable_pregunta,
        "clave": st.session_state.editable_clave,
        "justificacion_clave": st.session_state.editable_just_clave,
        "grafico_necesario_enunciado": st.session_state.editable_grafico_nec_enunciado,
        "opciones": {},
        "justificaciones_distractores": [
            {"opcion": "A", "justificacion": st.session_state.editable_just_a},
            {"opcion": "B", "justificacion": st.session_state.editable_just_b},
            {"opcion": "C", "justificacion": st.session_state.editable_just_c},
            {"opcion": "D", "justificacion": st.session_state.editable_just_d},
        ]
    }
    
    # Re-ensamble del gráfico del enunciado (lee el JSON que generamos o editamos)
    try:
        datos_editados["descripcion_grafico_enunciado"] = json.loads(st.session_state.editable_grafico_json_enunciado)
    except json.JSONDecodeError:
        st.error("El JSON del gráfico del enunciado tiene un error de formato, se guardará como texto.")
        datos_editados["descripcion_grafico_enunciado"] = st.session_state.editable_grafico_json_enunciado
    
    # Re-ensamble de las opciones (A, B, C, D) (lee el JSON que generamos o editamos)
    for letra in ["A", "B", "C", "D"]:
        opcion_data = {
            "texto": st.session_state[f"editable_opcion_{letra.lower()}_texto"],
            "grafico_necesario": st.session_state[f"editable_opcion_{letra.lower()}_grafico_nec"]
        }
        try:
            opcion_data["descripcion_grafico"] = json.loads(st.session_state[f"editable_opcion_{letra.lower()}_grafico_json"])
        except json.JSONDecodeError:
            opcion_data["descripcion_grafico"] = st.session_state[f"editable_opcion_{letra.lower()}_grafico_json"]
            st.error(f"El JSON del gráfico de la Opción {letra} tiene un error, se guardará como texto.")
        
        datos_editados["opciones"][letra] = opcion_data

    
    # --- GENERAR DATOS ADICIONALES ANTES DE LA DESCARGA ---
    taxonomia_actual = st.session_state.get('taxonomia_actual', {})
    
    with st.spinner("Generando oportunidad de mejora..."):
        # 1. Generar Oportunidad de Mejora (una sola vez)
        oportunidad_mejora = generar_oportunidad_mejora_llm(
            taxonomia_actual,
            datos_editados.get("justificacion_clave", "")
        )
    # --- FIN DE CAMBIOS ---

    col_word, col_excel = st.columns(2)
    
    with col_word:
        # 2. Pasar la oportunidad_mejora a la función de Word
        archivo_word = crear_word(datos_editados, taxonomia_actual, oportunidad_mejora)
        st.download_button(
            label="Descargar en Word (.docx)",
            data=archivo_word,
            file_name="item_espejo_auditado.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )
        
    with col_excel:
        # 3. Pasar los datos a la nueva función de Excel
        archivo_excel = crear_excel(
            datos_editados, 
            taxonomia_actual, 
            oportunidad_mejora
        )
        st.download_button(
            label="Descargar en Excel (.xlsx)",
            data=archivo_excel,
            file_name="item_espejo_auditado.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True
        )
